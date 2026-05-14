"""Парсер документов: PDF, DOCX, TXT, MD.

Возвращает текст. Большие документы обрезает до MAX_TEXT_CHARS,
чтобы не съесть контекст LLM.
"""
from __future__ import annotations

import io
import logging
from dataclasses import dataclass
from typing import Optional

log = logging.getLogger(__name__)

MAX_TEXT_CHARS = 12000


@dataclass
class ParsedDocument:
    filename: str
    extension: str
    text: str
    error: Optional[str] = None


def parse_document(filename: str, content: bytes) -> ParsedDocument:
    """Главная точка входа: определяет формат по расширению, парсит, возвращает текст."""
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""

    try:
        if ext in ("txt", "md", "csv", "log"):
            text = content.decode("utf-8", errors="replace")
        elif ext == "pdf":
            text = _parse_pdf(content)
        elif ext in ("docx", "doc"):
            text = _parse_docx(content)
        else:
            return ParsedDocument(
                filename=filename, extension=ext, text="",
                error=f"Формат .{ext} не поддерживается. Используй .pdf, .docx, .txt, .md.",
            )
    except Exception as exc:
        log.exception("Document parse failed: %s", filename)
        return ParsedDocument(filename=filename, extension=ext, text="", error=str(exc))

    if len(text) > MAX_TEXT_CHARS:
        text = text[:MAX_TEXT_CHARS] + "\n\n[...документ обрезан...]"
    return ParsedDocument(filename=filename, extension=ext, text=text)


def _parse_pdf(content: bytes) -> str:
    from pypdf import PdfReader
    reader = PdfReader(io.BytesIO(content))
    parts: list[str] = []
    for page in reader.pages:
        try:
            t = page.extract_text() or ""
        except Exception:
            t = ""
        if t.strip():
            parts.append(t)
    return "\n\n".join(parts)


def _parse_docx(content: bytes) -> str:
    from docx import Document
    doc = Document(io.BytesIO(content))
    parts: list[str] = []
    for p in doc.paragraphs:
        t = p.text.strip()
        if t:
            parts.append(t)
    # Также таблицы
    for tbl in doc.tables:
        for row in tbl.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)
